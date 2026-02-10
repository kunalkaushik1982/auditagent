"""
Annotator Engine for Word Documents.
File: backend/app/core/annotator.py

This module handles the injection of inline audits/comments into Word documents.
"""

import logging
import os
from typing import List, Dict, Any
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

logger = logging.getLogger(__name__)

class Annotator:
    """
    Handles the annotation of Word documents with audit findings.
    """
    
    def __init__(self):
        pass
        
    def annotate_document(self, original_path: str, findings: List[Dict[str, Any]]) -> str:
        """
        Open document, inject findings as highlights/comments, and save as new file.
        
        Args:
            original_path: Path to the original .docx file
            findings: List of finding dictionaries (must contain 'quote', 'description', 'severity')
            
        Returns:
            Path to the annotated document
        """
        if not os.path.exists(original_path):
            raise FileNotFoundError(f"Document not found: {original_path}")
            
        try:
            doc = Document(original_path)
            logger.info(f"Opened document for annotation: {original_path}")
            
            annotated_count = 0
            
            for finding in findings:
                quote = finding.get('quote')
                
                # If no quote or quote is "N/A", skip inline annotation (it will be in the report)
                if not quote or len(quote) < 5 or "N/A" in quote:
                    continue
                    
                # Search and Highlight
                if self._highlight_text(doc, quote, finding):
                    annotated_count += 1
            
            # Save properly
            base, ext = os.path.splitext(original_path)
            output_path = f"{base}_annotated{ext}"
            doc.save(output_path)
            
            logger.info(f"Saved annotated document: {output_path} ({annotated_count} annotations)")
            return output_path
            
        except Exception as e:
            logger.error(f"Error annotating document: {e}", exc_info=True)
            raise

    def annotate_pdf(self, original_path: str, findings: List[Dict[str, Any]]) -> str:
        """
        Annotate PDF with highlights based on findings.
        Includes robust search for text matching.
        """
        if not fitz:
            logger.warning("PyMuPDF (fitz) not installed. Skipping PDF annotation.")
            return original_path
            
        if not os.path.exists(original_path):
            raise FileNotFoundError(f"Document not found: {original_path}")
            
        try:
            doc = fitz.open(original_path)
            logger.info(f"Opened PDF for annotation: {original_path}")
            
            annotated_count = 0
            
            for finding in findings:
                quote = finding.get('quote')
                severity = (finding.get('severity') or 'low').lower()
                description = finding.get('description') or 'Issue found'
                
                if not quote or len(quote) < 5 or "N/A" in quote:
                    continue
                
                # Normalize quote (remove extra whitespace/newlines)
                clean_quote = " ".join(quote.split())
                
                # Search strategies
                found_instances = []
                
                # 1. Exact Search per page
                for page in doc:
                    # Try exact quote first
                    insts = page.search_for(quote)
                    
                    # Try cleaned quote if exact failed
                    if not insts and clean_quote != quote:
                        insts = page.search_for(clean_quote)
                        
                    # Try first 50 chars if long quote fails (fuzzy-ish)
                    if not insts and len(clean_quote) > 100:
                        short_quote = clean_quote[:50]
                        insts = page.search_for(short_quote)
                        if insts:
                            logger.info(f"Matched partial quote: '{short_quote}...'")
                    
                    if insts:
                        found_instances.append((page, insts))
                
                if found_instances:
                    annotated_count += 1
                    
                    # Set color
                    if severity == 'critical':
                        color = (1, 0, 0) # Red
                    elif severity == 'high':
                        color = (1, 0.5, 0) # Orange
                    elif severity == 'medium':
                        color = (1, 1, 0) # Yellow
                    else:
                        color = (0, 1, 0) # Green
                    
                    for page, insts in found_instances:
                        for inst in insts:
                            highlight = page.add_highlight_annot(inst)
                            highlight.set_colors(stroke=color)
                            highlight.set_info(content=description, title=f"Audit Finding ({severity.upper()})")
                            highlight.update()
                else:
                    logger.warning(f"Quote NOT found in PDF: '{quote[:50]}...' (Cleaned: '{clean_quote[:50]}...')")
            
            # Save output
            base, ext = os.path.splitext(original_path)
            output_path = f"{base}_annotated{ext}"
            doc.save(output_path)
            doc.close()
            
            logger.info(f"Saved annotated PDF: {output_path} ({annotated_count} findings highlighted)")
            return output_path
            
        except Exception as e:
            logger.error(f"Error annotating PDF: {e}", exc_info=True)
            raise

    def _highlight_text(self, doc, quote: str, finding: Dict[str, Any]) -> bool:
        """
        Search for quote in document and highlight it.
        Returns True if found and highlighted.
        """
        # Clean quotes (sometimes AI adds quotes around the text)
        clean_quote = quote.strip('"').strip("'").strip()
        
        for paragraph in doc.paragraphs:
            if clean_quote in paragraph.text:
                # Found the paragraph! 
                # Identifying exact run is hard, so for MVP we append a comment 
                # to the paragraph text directly in a distinct format.
                
                # 1. Add visual highlight to the paragraph (or try to find run)
                # For MVP: We will append the finding to the end of the paragraph in [RED]
                
                run = paragraph.add_run(f"\n[AUDIT {finding.get('finding_type', 'issue').upper()}: {finding.get('description', '')}]")
                run.bold = True
                
                # Color code based on severity/type
                severity = finding.get('severity', 'medium').lower()
                if severity == 'critical':
                    run.font.highlight_color = WD_COLOR_INDEX.RED
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif severity == 'high':
                    run.font.highlight_color = WD_COLOR_INDEX.DARK_RED
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif severity == 'medium':
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                
                return True
                
        return False
