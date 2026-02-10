from docx import Document
import os

# Create Project Plan (with issues)
doc = Document()
doc.add_heading('Simple Project Plan', 0)

doc.add_heading('Timeline', level=1)
doc.add_paragraph('The project will start on Jan 1st and end on Jan 15th. This 2 week timeline is fixed.')

doc.add_heading('Resources', level=1)
doc.add_paragraph('We have 1 developer assigned part-time.')

doc.add_heading('Budget', level=1)
doc.add_paragraph('The budget is $500.')

# Notice: Missing Risk Section
doc.save('sample_data/project_plan_v1.docx')
print('Created sample_data/project_plan_v1.docx')

# Create Checklist
checklist_doc = Document()
checklist_doc.add_heading('Project Plan Checklist', 0)
checklist_doc.add_paragraph('1. Is the project duration at least 1 month? (Timeline)')
checklist_doc.add_paragraph('2. Is there a dedicated Risk Management section?')
checklist_doc.add_paragraph('3. Is the budget sufficient for enterprise deployment?')

checklist_doc.save('sample_data/project_plan_checklist.docx')
print('Created sample_data/project_plan_checklist.docx')
